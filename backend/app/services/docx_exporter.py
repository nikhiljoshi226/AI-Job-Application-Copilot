import os
import uuid
from datetime import datetime
from typing import Dict, Any, Optional, Tuple
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.shared import OxmlElement, qn

from app.services.resume_renderer import ResumeRenderer
from app.core.config import settings


class DocxExporter:
    """
    Service to export tailored resumes to DOCX format using python-docx.
    """

    def __init__(self):
        self.renderer = ResumeRenderer()
        self.export_dir = Path(settings.UPLOAD_DIR) / "exports" / "docx"
        self.export_dir.mkdir(parents=True, exist_ok=True)
        
        # Template styles
        self.styles = {
            "name": {
                "size": Pt(16),
                "bold": True,
                "color": RGBColor(0, 0, 0),
                "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER
            },
            "section_title": {
                "size": Pt(12),
                "bold": True,
                "color": RGBColor(0, 0, 0),
                "space_after": Pt(12)
            },
            "company_title": {
                "size": Pt(11),
                "bold": True,
                "color": RGBColor(0, 0, 0),
                "space_after": Pt(6)
            },
            "normal": {
                "size": Pt(11),
                "color": RGBColor(0, 0, 0),
                "space_after": Pt(6)
            },
            "contact": {
                "size": Pt(10),
                "color": RGBColor(100, 100, 100),
                "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
                "space_after": Pt(12)
            },
            "bullet": {
                "size": Pt(11),
                "color": RGBColor(0, 0, 0),
                "space_after": Pt(3)
            },
            "date": {
                "size": Pt(10),
                "color": RGBColor(100, 100, 100),
                "italic": True,
                "space_after": Pt(6)
            }
        }

    def export_tailored_resume_to_docx(
        self,
        tailored_resume_data: Dict[str, Any],
        template: str = "professional",
        filename: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Export tailored resume to DOCX format.
        
        Args:
            tailored_resume_data: Tailored resume data with rendering data
            template: Template style to use
            filename: Optional custom filename
            
        Returns:
            Export result with file metadata
        """
        try:
            # Generate filename if not provided
            if not filename:
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                filename = f"tailored_resume_{timestamp}_{uuid.uuid4().hex[:8]}.docx"
            
            # Prepare rendering data
            rendering_data = self.renderer.prepare_for_docx_rendering(
                tailored_resume_data,
                template
            )
            
            # Create new document
            doc = Document()
            
            # Apply template-specific settings
            self._apply_document_settings(doc, template)
            
            # Build document sections
            self._build_header(doc, rendering_data["content"]["header"])
            self._build_summary(doc, rendering_data["content"]["summary"])
            self._build_experience(doc, rendering_data["content"]["experience"])
            self._build_education(doc, rendering_data["content"]["education"])
            self._build_skills(doc, rendering_data["content"]["skills"])
            self._build_projects(doc, rendering_data["content"]["projects"])
            self._build_certifications(doc, rendering_data["content"]["certifications"])
            self._build_languages(doc, rendering_data["content"]["languages"])
            
            # Save document
            file_path = self.export_dir / filename
            doc.save(str(file_path))
            
            # Generate file metadata
            file_metadata = {
                "filename": filename,
                "file_path": str(file_path),
                "file_size": file_path.stat().st_size,
                "template": template,
                "exported_at": datetime.utcnow().isoformat(),
                "content_hash": self._generate_content_hash(rendering_data),
                "download_url": f"/api/v1/exports/docx/download/{filename}"
            }
            
            return {
                "success": True,
                "file_metadata": file_metadata,
                "rendering_metadata": rendering_data.get("metadata", {}),
                "validation_errors": []
            }
            
        except Exception as e:
            return {
                "success": False,
                "error": str(e),
                "file_metadata": None,
                "rendering_metadata": None,
                "validation_errors": [f"Export failed: {str(e)}"]
            }

    def _apply_document_settings(self, doc: Document, template: str):
        """Apply document-wide settings based on template."""
        # Set margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

    def _build_header(self, doc: Document, header_data: Dict[str, Any]):
        """Build the header section with name and contact info."""
        if not header_data.get("name"):
            return
        
        # Name
        name_para = doc.add_paragraph()
        self._apply_paragraph_style(name_para, self.styles["name"])
        name_para.add_run(header_data["name"])
        
        # Contact information
        contact = header_data.get("contact", {})
        contact_parts = []
        
        if contact.get("email"):
            contact_parts.append(contact["email"])
        if contact.get("phone"):
            contact_parts.append(contact["phone"])
        if contact.get("location"):
            contact_parts.append(contact["location"])
        if contact.get("linkedin"):
            contact_parts.append(contact["linkedin"])
        if contact.get("github"):
            contact_parts.append(contact["github"])
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            self._apply_paragraph_style(contact_para, self.styles["contact"])
            contact_para.add_run(" | ".join(contact_parts))

    def _build_summary(self, doc: Document, summary_data: Dict[str, Any]):
        """Build the professional summary section."""
        if not summary_data.get("content"):
            return
        
        # Section title
        title_para = doc.add_paragraph()
        self._apply_paragraph_style(title_para, self.styles["section_title"])
        title_para.add_run(summary_data.get("title", "PROFESSIONAL SUMMARY"))
        
        # Summary content
        summary_para = doc.add_paragraph()
        self._apply_paragraph_style(summary_para, self.styles["normal"])
        summary_para.add_run(summary_data["content"])

    def _build_experience(self, doc: Document, experience_data: Dict[str, Any]):
        """Build the professional experience section."""
        entries = experience_data.get("entries", [])
        if not entries:
            return
        
        # Section title
        title_para = doc.add_paragraph()
        self._apply_paragraph_style(title_para, self.styles["section_title"])
        title_para.add_run(experience_data.get("title", "PROFESSIONAL EXPERIENCE"))
        
        # Experience entries
        for entry in entries:
            self._build_experience_entry(doc, entry)

    def _build_experience_entry(self, doc: Document, entry: Dict[str, Any]):
        """Build a single experience entry."""
        # Title and company
        title_company = f"{entry.get('title', '')} at {entry.get('company', '')}"
        title_para = doc.add_paragraph()
        self._apply_paragraph_style(title_para, self.styles["company_title"])
        title_para.add_run(title_company)
        
        # Dates
        dates = entry.get("dates", "")
        if dates:
            date_para = doc.add_paragraph()
            self._apply_paragraph_style(date_para, self.styles["date"])
            date_para.add_run(dates)
        
        # Description
        description = entry.get("description", "")
        if description:
            desc_para = doc.add_paragraph()
            self._apply_paragraph_style(desc_para, self.styles["normal"])
            desc_para.add_run(description)
        
        # Achievements (bullet points)
        achievements = entry.get("achievements", [])
        if achievements:
            for achievement in achievements:
                bullet_para = doc.add_paragraph()
                self._apply_paragraph_style(bullet_para, self.styles["bullet"])
                self._add_bullet_point(bullet_para, achievement)
        
        # Add spacing between entries
        doc.add_paragraph().add_run("")

    def _build_education(self, doc: Document, education_data: Dict[str, Any]):
        """Build the education section."""
        entries = education_data.get("entries", [])
        if not entries:
            return
        
        # Section title
        title_para = doc.add_paragraph()
        self._apply_paragraph_style(title_para, self.styles["section_title"])
        title_para.add_run(education_data.get("title", "EDUCATION"))
        
        # Education entries
        for entry in entries:
            self._build_education_entry(doc, entry)

    def _build_education_entry(self, doc: Document, entry: Dict[str, Any]):
        """Build a single education entry."""
        # Degree and major
        degree_major = f"{entry.get('degree', '')} in {entry.get('major', '')}"
        degree_para = doc.add_paragraph()
        self._apply_paragraph_style(degree_para, self.styles["company_title"])
        degree_para.add_run(degree_major)
        
        # University and location
        university = entry.get("university", "")
        location = entry.get("location", "")
        if location:
            university += f", {location}"
        
        if university:
            uni_para = doc.add_paragraph()
            self._apply_paragraph_style(uni_para, self.styles["normal"])
            uni_para.add_run(university)
        
        # Graduation year
        grad_year = entry.get("graduation_year", "")
        if grad_year:
            year_para = doc.add_paragraph()
            self._apply_paragraph_style(year_para, self.styles["date"])
            year_para.add_run(f"Graduated: {grad_year}")
        
        # GPA
        gpa = entry.get("gpa", "")
        if gpa:
            gpa_para = doc.add_paragraph()
            self._apply_paragraph_style(gpa_para, self.styles["normal"])
            gpa_para.add_run(f"GPA: {gpa}")
        
        # Add spacing
        doc.add_paragraph().add_run("")

    def _build_skills(self, doc: Document, skills_data: Dict[str, Any]):
        """Build the skills section."""
        categories = skills_data.get("categories", [])
        if not categories:
            return
        
        # Section title
        title_para = doc.add_paragraph()
        self._apply_paragraph_style(title_para, self.styles["section_title"])
        title_para.add_run(skills_data.get("title", "SKILLS"))
        
        # Skill categories
        for category in categories:
            self._build_skill_category(doc, category)

    def _build_skill_category(self, doc: Document, category: Dict[str, Any]):
        """Build a single skill category."""
        category_name = category.get("name", "")
        skills = category.get("skills", [])
        
        if not skills:
            return
        
        # Category name
        cat_para = doc.add_paragraph()
        self._apply_paragraph_style(cat_para, self.styles["company_title"])
        cat_para.add_run(f"{category_name}:")
        
        # Skills list
        skill_names = []
        for skill in skills:
            skill_name = skill.get("name", "")
            if skill_name:
                if skill.get("highlight", False):
                    skill_name += " *"  # Mark added skills
                skill_names.append(skill_name)
        
        if skill_names:
            skills_para = doc.add_paragraph()
            self._apply_paragraph_style(skills_para, self.styles["normal"])
            skills_para.add_run(", ".join(skill_names))
        
        # Add spacing
        doc.add_paragraph().add_run("")

    def _build_projects(self, doc: Document, projects_data: Optional[Dict[str, Any]]):
        """Build the projects section."""
        if not projects_data:
            return
        
        entries = projects_data.get("entries", [])
        if not entries:
            return
        
        # Section title
        title_para = doc.add_paragraph()
        self._apply_paragraph_style(title_para, self.styles["section_title"])
        title_para.add_run(projects_data.get("title", "PROJECTS"))
        
        # Limit to max entries
        max_entries = projects_data.get("max_entries", 3)
        limited_entries = entries[:max_entries]
        
        # Project entries
        for entry in limited_entries:
            self._build_project_entry(doc, entry)

    def _build_project_entry(self, doc: Document, entry: Dict[str, Any]):
        """Build a single project entry."""
        # Project name
        name_para = doc.add_paragraph()
        self._apply_paragraph_style(name_para, self.styles["company_title"])
        name_para.add_run(entry.get("name", ""))
        
        # Dates
        dates = entry.get("dates", "")
        if dates:
            date_para = doc.add_paragraph()
            self._apply_paragraph_style(date_para, self.styles["date"])
            date_para.add_run(dates)
        
        # Description
        description = entry.get("description", "")
        if description:
            desc_para = doc.add_paragraph()
            self._apply_paragraph_style(desc_para, self.styles["normal"])
            desc_para.add_run(description)
        
        # Technologies
        technologies = entry.get("technologies", [])
        if technologies:
            tech_para = doc.add_paragraph()
            self._apply_paragraph_style(tech_para, self.styles["normal"])
            tech_para.add_run(f"Technologies: {', '.join(technologies)}")
        
        # Add spacing
        doc.add_paragraph().add_run("")

    def _build_certifications(self, doc: Document, certifications_data: Optional[Dict[str, Any]]):
        """Build the certifications section."""
        if not certifications_data:
            return
        
        entries = certifications_data.get("entries", [])
        if not entries:
            return
        
        # Section title
        title_para = doc.add_paragraph()
        self._apply_paragraph_style(title_para, self.styles["section_title"])
        title_para.add_run(certifications_data.get("title", "CERTIFICATIONS"))
        
        # Certification entries
        for entry in entries:
            self._build_certification_entry(doc, entry)

    def _build_certification_entry(self, doc: Document, entry: Dict[str, Any]):
        """Build a single certification entry."""
        # Certification name and issuer
        name = entry.get("name", "")
        issuer = entry.get("issuer", "")
        if issuer:
            cert_text = f"{name} - {issuer}"
        else:
            cert_text = name
        
        cert_para = doc.add_paragraph()
        self._apply_paragraph_style(cert_para, self.styles["normal"])
        cert_para.add_run(cert_text)
        
        # Date
        cert_date = entry.get("date", "") or entry.get("year", "")
        if cert_date:
            date_para = doc.add_paragraph()
            self._apply_paragraph_style(date_para, self.styles["date"])
            date_para.add_run(cert_date)
        
        # Add spacing
        doc.add_paragraph().add_run("")

    def _build_languages(self, doc: Document, languages_data: Optional[Dict[str, Any]]):
        """Build the languages section."""
        if not languages_data:
            return
        
        languages = languages_data.get("languages", [])
        if not languages:
            return
        
        # Section title
        title_para = doc.add_paragraph()
        self._apply_paragraph_style(title_para, self.styles["section_title"])
        title_para.add_run(languages_data.get("title", "LANGUAGES"))
        
        # Languages list
        language_texts = []
        for lang in languages:
            lang_name = lang.get("language", "")
            proficiency = lang.get("proficiency", "")
            if proficiency:
                lang_text = f"{lang_name} ({proficiency})"
            else:
                lang_text = lang_name
            language_texts.append(lang_text)
        
        if language_texts:
            lang_para = doc.add_paragraph()
            self._apply_paragraph_style(lang_para, self.styles["normal"])
            lang_para.add_run(", ".join(language_texts))

    def _apply_paragraph_style(self, paragraph, style: Dict[str, Any]):
        """Apply style to a paragraph."""
        # Alignment
        if "alignment" in style:
            paragraph.alignment = style["alignment"]
        
        # Spacing
        if "space_after" in style:
            paragraph.paragraph_format.space_after = style["space_after"]
        
        # Apply to all runs in paragraph
        for run in paragraph.runs:
            self._apply_run_style(run, style)

    def _apply_run_style(self, run, style: Dict[str, Any]):
        """Apply style to a text run."""
        if "size" in style:
            run.font.size = style["size"]
        
        if "bold" in style:
            run.font.bold = style["bold"]
        
        if "italic" in style:
            run.font.italic = style["italic"]
        
        if "color" in style:
            run.font.color.rgb = style["color"]

    def _add_bullet_point(self, paragraph, text: str):
        """Add a bullet point to a paragraph."""
        # Create bullet paragraph
        run = paragraph.add_run("• ")
        run.font.bold = True
        
        # Add text
        text_run = paragraph.add_run(text)
        text_run.font.size = Pt(11)

    def _generate_content_hash(self, rendering_data: Dict[str, Any]) -> str:
        """Generate a hash of the content for change detection."""
        import hashlib
        content_str = str(rendering_data.get("content", {}))
        return hashlib.md5(content_str.encode()).hexdigest()

    def validate_export_data(self, tailored_resume_data: Dict[str, Any]) -> Tuple[bool, list]:
        """
        Validate the data before export.
        
        Args:
            tailored_resume_data: Tailored resume data
            
        Returns:
            Tuple of (is_valid, list_of_validation_errors)
        """
        errors = []
        
        # Check if rendering data exists
        if "rendering_data" not in tailored_resume_data:
            errors.append("Missing rendering data in tailored resume")
            return False, errors
        
        rendering_data = tailored_resume_data["rendering_data"]
        
        # Validate required sections
        content = rendering_data.get("content", {})
        
        if not content.get("header", {}).get("name"):
            errors.append("Missing name in header")
        
        if not content.get("summary", {}).get("content"):
            errors.append("Missing summary content")
        
        # Check experience entries
        experience = content.get("experience", {}).get("entries", [])
        if not experience:
            errors.append("No experience entries found")
        
        for i, entry in enumerate(experience):
            if not entry.get("title"):
                errors.append(f"Experience entry {i+1} missing title")
            if not entry.get("company"):
                errors.append(f"Experience entry {i+1} missing company")
        
        return len(errors) == 0, errors

    def get_export_statistics(self, file_path: str) -> Dict[str, Any]:
        """
        Get statistics about the exported file.
        
        Args:
            file_path: Path to the exported file
            
        Returns:
            File statistics
        """
        try:
            path = Path(file_path)
            if not path.exists():
                return {"error": "File not found"}
            
            stat = path.stat()
            
            # Try to get word count (basic estimation)
            try:
                doc = Document(str(path))
                word_count = sum(len(p.text.split()) for p in doc.paragraphs)
                paragraph_count = len(doc.paragraphs)
            except:
                word_count = 0
                paragraph_count = 0
            
            return {
                "file_size": stat.st_size,
                "file_size_human": self._format_file_size(stat.st_size),
                "created_at": datetime.fromtimestamp(stat.st_ctime).isoformat(),
                "modified_at": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                "word_count": word_count,
                "paragraph_count": paragraph_count,
                "file_extension": path.suffix,
                "file_name": path.name
            }
            
        except Exception as e:
            return {"error": f"Error getting statistics: {str(e)}"}

    def _format_file_size(self, size_bytes: int) -> str:
        """Format file size in human readable format."""
        if size_bytes < 1024:
            return f"{size_bytes} B"
        elif size_bytes < 1024 * 1024:
            return f"{size_bytes / 1024:.1f} KB"
        else:
            return f"{size_bytes / (1024 * 1024):.1f} MB"

    def cleanup_old_exports(self, days_old: int = 30):
        """
        Clean up old export files.
        
        Args:
            days_old: Age in days to delete files
        """
        try:
            cutoff_time = datetime.now().timestamp() - (days_old * 24 * 3600)
            
            for file_path in self.export_dir.glob("*.docx"):
                if file_path.stat().st_mtime < cutoff_time:
                    file_path.unlink()
                    
        except Exception as e:
            print(f"Error cleaning up old exports: {e}")

import pytest
import tempfile
import shutil
from pathlib import Path
from unittest.mock import patch, MagicMock
from datetime import datetime

from app.services.docx_exporter import DocxExporter


class TestDocxExporter:
    
    def setup_method(self):
        """Setup test fixtures."""
        # Create temporary directory for exports
        self.temp_dir = tempfile.mkdtemp()
        self.exporter = DocxExporter()
        self.exporter.export_dir = Path(self.temp_dir)
        
        # Sample tailored resume data
        self.tailored_resume_data = {
            "rendering_data": {
                "header": {
                    "name": "John Doe",
                    "contact": {
                        "email": "john@example.com",
                        "phone": "555-1234",
                        "location": "New York, NY",
                        "linkedin": "linkedin.com/in/johndoe",
                        "github": "github.com/johndoe"
                    }
                },
                "summary": {
                    "title": "Professional Summary",
                    "content": "Senior Software Engineer with 5 years of experience in full-stack development"
                },
                "skills": {
                    "title": "Skills",
                    "categories": [
                        {
                            "name": "Technical Skills",
                            "skills": [
                                {"name": "Python", "level": "5 years", "highlight": False},
                                {"name": "Django", "level": "3 years", "highlight": True}
                            ]
                        },
                        {
                            "name": "Soft Skills",
                            "skills": [
                                {"name": "Communication", "level": "Strong", "highlight": False}
                            ]
                        }
                    ]
                },
                "experience": {
                    "title": "Professional Experience",
                    "entries": [
                        {
                            "title": "Software Engineer",
                            "company": "Tech Corp",
                            "location": "San Francisco, CA",
                            "dates": "January 2021 - March 2023",
                            "description": "Developed scalable web applications",
                            "achievements": [
                                "Improved application performance by 40%",
                                "Led team of 3 junior developers"
                            ],
                            "technologies": ["Python", "Django", "React"]
                        }
                    ]
                },
                "education": {
                    "title": "Education",
                    "entries": [
                        {
                            "degree": "Bachelor of Science",
                            "major": "Computer Science",
                            "university": "State University",
                            "location": "Boston, MA",
                            "graduation_year": "2020",
                            "gpa": "3.8"
                        }
                    ]
                },
                "projects": {
                    "title": "Projects",
                    "entries": [
                        {
                            "name": "E-commerce Platform",
                            "description": "Built a scalable e-commerce web application",
                            "technologies": ["React", "Node.js", "MongoDB"],
                            "dates": "June 2022 - December 2022"
                        }
                    ],
                    "max_entries": 3
                },
                "certifications": {
                    "title": "Certifications",
                    "entries": [
                        {
                            "name": "AWS Certified Developer",
                            "issuer": "Amazon Web Services",
                            "date": "2022",
                            "credential_id": "AWS-DEV-123"
                        }
                    ]
                },
                "languages": {
                    "title": "Languages",
                    "languages": [
                        {"language": "English", "proficiency": "Native"},
                        {"language": "Spanish", "proficiency": "Conversational"}
                    ]
                }
            },
            "metadata": {
                "final_alignment_score": 85,
                "truthfulness_score": 0.88,
                "processing_time_ms": 1250,
                "created_at": "2023-01-01T00:00:00Z"
            }
        }

    def teardown_method(self):
        """Cleanup test fixtures."""
        shutil.rmtree(self.temp_dir, ignore_errors=True)

    def test_export_tailored_resume_to_docx_success(self):
        """Test successful DOCX export."""
        result = self.exporter.export_tailored_resume_to_docx(
            self.tailored_resume_data,
            template="professional",
            filename="test_resume.docx"
        )
        
        # Check export result
        assert result["success"] == True
        assert result["file_metadata"] is not None
        assert result["validation_errors"] == []
        
        # Check file metadata
        metadata = result["file_metadata"]
        assert metadata["filename"] == "test_resume.docx"
        assert metadata["template"] == "professional"
        assert metadata["file_size"] > 0
        assert "exported_at" in metadata
        assert "download_url" in metadata
        
        # Check file was created
        file_path = Path(self.temp_dir) / "test_resume.docx"
        assert file_path.exists()
        assert file_path.stat().st_size > 0

    def test_export_tailored_resume_to_docx_auto_filename(self):
        """Test DOCX export with auto-generated filename."""
        result = self.exporter.export_tailored_resume_to_docx(
            self.tailored_resume_data,
            template="modern"
        )
        
        assert result["success"] == True
        assert result["file_metadata"] is not None
        
        # Check filename format
        metadata = result["file_metadata"]
        assert metadata["filename"].startswith("tailored_resume_")
        assert metadata["filename"].endswith(".docx")
        assert metadata["template"] == "modern"

    def test_export_tailored_resume_to_docx_validation_error(self):
        """Test DOCX export with validation errors."""
        # Create invalid data (missing required sections)
        invalid_data = {
            "rendering_data": {
                "header": {},  # Missing name
                "summary": {},  # Missing content
                "experience": {"entries": []}  # No experience entries
            }
        }
        
        result = self.exporter.export_tailored_resume_to_docx(invalid_data)
        
        assert result["success"] == False
        assert result["error"] == "Validation failed"
        assert len(result["validation_errors"]) > 0
        assert any("Missing name" in error for error in result["validation_errors"])
        assert any("Missing summary" in error for error in result["validation_errors"])
        assert any("No experience entries" in error for error in result["validation_errors"])

    def test_export_tailored_resume_to_docx_missing_rendering_data(self):
        """Test DOCX export with missing rendering data."""
        invalid_data = {}
        
        result = self.exporter.export_tailored_resume_to_docx(invalid_data)
        
        assert result["success"] == False
        assert result["error"] == "Validation failed"
        assert len(result["validation_errors"]) > 0
        assert "Missing rendering data" in result["validation_errors"][0]

    @patch('app.services.docx_exporter.Document')
    def test_export_tailored_resume_to_docx_document_error(self, mock_document):
        """Test DOCX export with document creation error."""
        # Mock Document to raise an exception
        mock_document.side_effect = Exception("Document creation failed")
        
        result = self.exporter.export_tailored_resume_to_docx(self.tailored_resume_data)
        
        assert result["success"] == False
        assert "Export failed" in result["error"]
        assert result["file_metadata"] is None

    def test_validate_export_data_valid(self):
        """Test validation of valid export data."""
        is_valid, errors = self.exporter.validate_export_data(self.tailored_resume_data)
        
        assert is_valid == True
        assert len(errors) == 0

    def test_validate_export_data_invalid(self):
        """Test validation of invalid export data."""
        invalid_data = {
            "rendering_data": {
                "header": {},  # Missing name
                "summary": {},  # Missing content
                "experience": {"entries": []}  # No experience entries
            }
        }
        
        is_valid, errors = self.exporter.validate_export_data(invalid_data)
        
        assert is_valid == False
        assert len(errors) >= 3  # Should catch all three issues

    def test_validate_export_data_missing_rendering_data(self):
        """Test validation with missing rendering data."""
        invalid_data = {}
        
        is_valid, errors = self.exporter.validate_export_data(invalid_data)
        
        assert is_valid == False
        assert len(errors) == 1
        assert "Missing rendering data" in errors[0]

    def test_get_export_statistics_existing_file(self):
        """Test getting statistics for existing file."""
        # Create a test file
        test_file = self.exporter.export_dir / "test_stats.docx"
        test_file.write_text("dummy content")  # Write some content
        
        statistics = self.exporter.get_export_statistics(str(test_file))
        
        assert "file_size" in statistics
        assert "file_size_human" in statistics
        assert "created_at" in statistics
        assert "modified_at" in statistics
        assert "file_extension" in statistics
        assert "file_name" in statistics
        assert statistics["file_name"] == "test_stats.docx"
        assert statistics["file_extension"] == ".docx"

    def test_get_export_statistics_nonexistent_file(self):
        """Test getting statistics for non-existent file."""
        statistics = self.exporter.get_export_statistics("nonexistent.docx")
        
        assert "error" in statistics
        assert "File not found" in statistics["error"]

    def test_format_file_size(self):
        """Test file size formatting."""
        # Test bytes
        assert self.exporter._format_file_size(512) == "512 B"
        
        # Test kilobytes
        assert self.exporter._format_file_size(1536) == "1.5 KB"
        
        # Test megabytes
        assert self.exporter._format_file_size(2 * 1024 * 1024) == "2.0 MB"

    def test_generate_content_hash(self):
        """Test content hash generation."""
        rendering_data = {
            "content": {
                "header": {"name": "John Doe"},
                "summary": {"content": "Software Engineer"}
            }
        }
        
        hash1 = self.exporter._generate_content_hash(rendering_data)
        hash2 = self.exporter._generate_content_hash(rendering_data)
        
        # Same content should generate same hash
        assert hash1 == hash2
        assert len(hash1) == 32  # MD5 hash length

    def test_cleanup_old_exports(self):
        """Test cleanup of old export files."""
        # Create some test files with different timestamps
        old_file = self.exporter.export_dir / "old_export.docx"
        new_file = self.exporter.export_dir / "new_export.docx"
        
        old_file.write_text("old content")
        new_file.write_text("new content")
        
        # Mock file timestamps
        old_time = datetime.now().timestamp() - (35 * 24 * 3600)  # 35 days ago
        new_time = datetime.now().timestamp() - (5 * 24 * 3600)   # 5 days ago
        
        import os
        import time
        
        os.utime(old_file, (old_time, old_time))
        os.utime(new_file, (new_time, new_time))
        
        # Cleanup files older than 30 days
        self.exporter.cleanup_old_exports(days_old=30)
        
        # Old file should be deleted, new file should remain
        assert not old_file.exists()
        assert new_file.exists()

    def test_build_header_complete(self):
        """Test building complete header section."""
        from docx import Document
        
        doc = Document()
        header_data = self.tailored_resume_data["rendering_data"]["header"]
        
        self.exporter._build_header(doc, header_data)
        
        # Should have name paragraph and contact paragraph
        assert len(doc.paragraphs) >= 2
        
        # Check name is present
        name_found = False
        for para in doc.paragraphs:
            if "John Doe" in para.text:
                name_found = True
                break
        assert name_found

    def test_build_summary_complete(self):
        """Test building complete summary section."""
        from docx import Document
        
        doc = Document()
        summary_data = self.tailored_resume_data["rendering_data"]["summary"]
        
        self.exporter._build_summary(doc, summary_data)
        
        # Should have title and content paragraphs
        assert len(doc.paragraphs) >= 2
        
        # Check content is present
        content_found = False
        for para in doc.paragraphs:
            if "Senior Software Engineer" in para.text:
                content_found = True
                break
        assert content_found

    def test_build_experience_complete(self):
        """Test building complete experience section."""
        from docx import Document
        
        doc = Document()
        experience_data = self.tailored_resume_data["rendering_data"]["experience"]
        
        self.exporter._build_experience(doc, experience_data)
        
        # Should have title and at least one entry
        assert len(doc.paragraphs) >= 3
        
        # Check company and title are present
        text_content = "\n".join(para.text for para in doc.paragraphs)
        assert "Software Engineer" in text_content
        assert "Tech Corp" in text_content

    def test_build_skills_complete(self):
        """Test building complete skills section."""
        from docx import Document
        
        doc = Document()
        skills_data = self.tailored_resume_data["rendering_data"]["skills"]
        
        self.exporter._build_skills(doc, skills_data)
        
        # Should have title and category sections
        assert len(doc.paragraphs) >= 4
        
        # Check skills are present
        text_content = "\n".join(para.text for para in doc.paragraphs)
        assert "Technical Skills" in text_content
        assert "Python" in text_content
        assert "Django" in text_content

    def test_apply_paragraph_style(self):
        """Test applying paragraph styles."""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
        
        doc = Document()
        para = doc.add_paragraph("Test text")
        
        style = {
            "size": Pt(12),
            "bold": True,
            "alignment": WD_PARAGRAPH_ALIGNMENT.CENTER,
            "space_after": Pt(6)
        }
        
        self.exporter._apply_paragraph_style(para, style)
        
        # Check style was applied (basic checks)
        assert para.alignment == WD_PARAGRAPH_ALIGNMENT.CENTER
        assert para.paragraph_format.space_after == Pt(6)

    def test_apply_run_style(self):
        """Test applying run styles."""
        from docx import Document
        from docx.shared import Pt, RGBColor
        
        doc = Document()
        para = doc.add_paragraph()
        run = para.add_run("Test text")
        
        style = {
            "size": Pt(14),
            "bold": True,
            "color": RGBColor(255, 0, 0)
        }
        
        self.exporter._apply_run_style(run, style)
        
        # Check style was applied
        assert run.font.size == Pt(14)
        assert run.font.bold == True

    def test_add_bullet_point(self):
        """Test adding bullet points."""
        from docx import Document
        
        doc = Document()
        para = doc.add_paragraph()
        
        self.exporter._add_bullet_point(para, "Test bullet point")
        
        # Check bullet was added
        assert "•" in para.text
        assert "Test bullet point" in para.text

    def test_build_education_entry(self):
        """Test building education entry."""
        from docx import Document
        
        doc = Document()
        education_entry = {
            "degree": "Bachelor of Science",
            "major": "Computer Science",
            "university": "State University",
            "location": "Boston, MA",
            "graduation_year": "2020",
            "gpa": "3.8"
        }
        
        self.exporter._build_education_entry(doc, education_entry)
        
        # Check education content is present
        text_content = "\n".join(para.text for para in doc.paragraphs)
        assert "Bachelor of Science" in text_content
        assert "Computer Science" in text_content
        assert "State University" in text_content

    def test_build_project_entry(self):
        """Test building project entry."""
        from docx import Document
        
        doc = Document()
        project_entry = {
            "name": "E-commerce Platform",
            "description": "Built a scalable e-commerce web application",
            "technologies": ["React", "Node.js", "MongoDB"],
            "dates": "June 2022 - December 2022"
        }
        
        self.exporter._build_project_entry(doc, project_entry)
        
        # Check project content is present
        text_content = "\n".join(para.text for para in doc.paragraphs)
        assert "E-commerce Platform" in text_content
        assert "React" in text_content

    def test_export_with_different_templates(self):
        """Test export with different templates."""
        templates = ["professional", "modern", "technical", "creative"]
        
        for template in templates:
            result = self.exporter.export_tailored_resume_to_docx(
                self.tailored_resume_data,
                template=template,
                filename=f"test_{template}.docx"
            )
            
            assert result["success"] == True
            assert result["file_metadata"]["template"] == template
            
            # Check file was created
            file_path = Path(self.temp_dir) / f"test_{template}.docx"
            assert file_path.exists()

    def test_export_with_minimal_data(self):
        """Test export with minimal but valid data."""
        minimal_data = {
            "rendering_data": {
                "header": {
                    "name": "Jane Doe",
                    "contact": {"email": "jane@example.com"}
                },
                "summary": {
                    "title": "Summary",
                    "content": "Experienced professional"
                },
                "experience": {
                    "title": "Experience",
                    "entries": [
                        {
                            "title": "Manager",
                            "company": "ABC Corp",
                            "dates": "2020-2023"
                        }
                    ]
                },
                "skills": {"categories": []},
                "education": {"entries": []},
                "projects": None,
                "certifications": None,
                "languages": None
            }
        }
        
        result = self.exporter.export_tailored_resume_to_docx(minimal_data)
        
        assert result["success"] == True
        assert result["file_metadata"] is not None


if __name__ == "__main__":
    pytest.main([__file__])
